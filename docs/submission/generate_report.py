#!/usr/bin/env python3
"""
Generate report_final.docx from report_template.docx + report_draft.md.

Fills the cover table, replaces placeholder text in each section with
the markdown draft content, converts markdown tables to Word tables,
and formats code blocks with monospace font.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE_DIR = Path(__file__).parent
TEMPLATE = BASE_DIR / "report_template.docx"
DRAFT = BASE_DIR / "report_draft.md"
OUTPUT = BASE_DIR / "report_final.docx"

# ---------------------------------------------------------------------------
# Parse the markdown
# ---------------------------------------------------------------------------

def parse_markdown(path: Path) -> dict:
    """Return a dict with cover info and section content."""
    text = path.read_text(encoding="utf-8")
    lines = text.split("\n")

    # Cover info
    cover = {}
    cover_keys = [
        "応募代表者氏名", "チーム名", "選択課題ID",
        "選択課題名", "提案名",
    ]
    for line in lines:
        for key in cover_keys:
            if line.strip().startswith(f"- {key}:"):
                cover[key] = line.split(":", 1)[1].strip()

    # Split into sections by ## headings
    sections = {}
    current_key = None
    current_lines = []

    for line in lines:
        m = re.match(r"^##\s+(.+)", line)
        if m:
            if current_key is not None:
                sections[current_key] = "\n".join(current_lines).strip()
            current_key = m.group(1).strip()
            current_lines = []
        elif current_key is not None:
            current_lines.append(line)

    if current_key is not None:
        sections[current_key] = "\n".join(current_lines).strip()

    return {"cover": cover, "sections": sections}


def extract_subsection(section_text: str, subsection_prefix: str) -> str:
    """Extract content under a ### subsection heading."""
    lines = section_text.split("\n")
    collecting = False
    result = []
    for line in lines:
        if re.match(r"^###\s+" + re.escape(subsection_prefix), line):
            collecting = True
            continue
        elif re.match(r"^###\s+", line):
            if collecting:
                break
        elif collecting:
            result.append(line)
    return "\n".join(result).strip()


def get_section_body(section_text: str, include_subsections: bool = False) -> str:
    """Get the body text of a section.
    If include_subsections is True, return the full section text.
    Otherwise, return only text before the first ### heading."""
    cleaned = re.sub(r"^---\s*$", "", section_text, flags=re.MULTILINE).strip()
    if include_subsections or "###" not in cleaned:
        return cleaned
    # Return only the text before the first ###
    lines = cleaned.split("\n")
    result = []
    for line in lines:
        if re.match(r"^###\s+", line):
            break
        result.append(line)
    return "\n".join(result).strip()


# ---------------------------------------------------------------------------
# Word document helpers
# ---------------------------------------------------------------------------

def clear_paragraph(para):
    """Remove all runs from a paragraph."""
    for run in para.runs:
        run._element.getparent().remove(run._element)
    # Also clear any remaining text nodes
    el = para._element
    for child in list(el):
        if child.tag.endswith("}r"):
            el.remove(child)


def set_paragraph_text(para, text, bold=False, font_name=None, font_size=None):
    """Set paragraph text, clearing existing content."""
    clear_paragraph(para)
    run = para.add_run(text)
    if bold:
        run.font.bold = True
    if font_name:
        run.font.name = font_name
        # For CJK fonts, also set the East Asian font
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = run._element.makeelement(qn("w:rFonts"), {})
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), font_name)
    if font_size:
        run.font.size = font_size


def add_paragraph_after(doc, ref_para, text="", style=None):
    """Insert a new paragraph after ref_para in the document body."""
    new_para = doc.add_paragraph(text, style=style)
    # Move the new paragraph element right after ref_para
    ref_para._element.addnext(new_para._element)
    return new_para


def add_formatted_content(doc, ref_para, content_text):
    """Parse markdown content and add formatted paragraphs after ref_para.
    Returns the last paragraph added (for chaining)."""
    lines = content_text.split("\n")
    last_para = ref_para
    i = 0

    while i < len(lines):
        line = lines[i]

        # Code block
        if line.strip().startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```

            code_text = "\n".join(code_lines)
            p = add_paragraph_after(doc, last_para)
            run = p.add_run(code_text)
            run.font.name = "Courier New"
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = run._element.makeelement(qn("w:rFonts"), {})
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:eastAsia"), "Courier New")
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            # Add light gray shading
            pPr = p._element.get_or_add_pPr()
            shd = p._element.makeelement(qn("w:shd"), {
                qn("w:val"): "clear",
                qn("w:color"): "auto",
                qn("w:fill"): "F2F2F2",
            })
            pPr.append(shd)
            last_para = p
            continue

        # Markdown table
        if line.strip().startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1

            # Parse table
            header_cells = [c.strip() for c in table_lines[0].strip().strip("|").split("|")]
            # Skip separator line (|---|---|...)
            data_rows = []
            for tl in table_lines[2:]:
                row_cells = [c.strip() for c in tl.strip().strip("|").split("|")]
                data_rows.append(row_cells)

            ncols = len(header_cells)
            nrows = len(data_rows) + 1  # +1 for header

            # Create table in the document
            table = doc.add_table(rows=nrows, cols=ncols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"

            # Fill header
            for ci, val in enumerate(header_cells):
                cell = table.rows[0].cells[ci]
                cell.text = val
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
                    run.font.size = Pt(9)

            # Fill data
            for ri, row_data in enumerate(data_rows):
                for ci in range(min(len(row_data), ncols)):
                    cell = table.rows[ri + 1].cells[ci]
                    cell.text = row_data[ci]
                    for run in cell.paragraphs[0].runs:
                        run.font.size = Pt(9)

            # Move the table element after last_para
            last_para._element.addnext(table._element)

            # Add an empty paragraph after the table for spacing
            spacer = add_paragraph_after(doc, last_para)
            # Move table between last_para and spacer
            spacer._element.addprevious(table._element)
            last_para = spacer
            continue

        # Regular text line
        if line.strip() == "---":
            i += 1
            continue

        if line.strip() == "":
            # Skip empty lines but don't create empty paragraphs
            i += 1
            continue

        # Bold handling: **text** -> bold run
        p = add_paragraph_after(doc, last_para)
        _add_rich_text(p, line)
        last_para = p
        i += 1

    return last_para


def _add_rich_text(para, text):
    """Add text to a paragraph, handling **bold** and plain text."""
    # Split on **...**
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.font.bold = True
        else:
            if part:
                para.add_run(part)


# ---------------------------------------------------------------------------
# Map template paragraphs to section content
# ---------------------------------------------------------------------------

# Template paragraph index -> what content goes there
# Based on the template structure analyzed above:
# [16] "1. 要約"           -> heading (keep)
# [17] placeholder         -> replace with section 1 body
# [19] "2. 課題の内容..."   -> heading (keep)
# [20] "2.1 課題の背景"    -> sub-heading (keep)
# [21] placeholder         -> replace with 2.1 content
# etc.

SECTION_MAP = {
    # (paragraph_index_of_placeholder, section_key, subsection_prefix_or_None)
    # subsection=None -> get body before first ###
    # subsection="X.Y" -> get content under ### X.Y heading
    # subsection="__FULL__" -> get entire section including all subsections
    17: ("1. 要約（1000文字以内）", None),
    21: ("2. 課題の内容・研究の目的", "2.1"),
    24: ("2. 課題の内容・研究の目的", "2.2"),
    27: ("2. 課題の内容・研究の目的", "2.3"),
    30: ("2. 課題の内容・研究の目的", "2.4"),
    34: ("3. 解決案の内容", "3.1"),
    37: ("3. 解決案の内容", "3.2"),
    40: ("3. 解決案の内容", "3.3"),
    43: ("3. 解決案の内容", "3.4"),
    46: ("3. 解決案の内容", "3.5"),
    49: ("4. 研究結果", "__FULL__"),
    53: ("5. 考察", "5.1"),
    56: ("5. 考察", "5.2"),
    59: ("5. 考察", "5.3"),
    62: ("5. 考察", "5.4"),
    66: ("6. まとめ", "6.1"),
    69: ("6. まとめ", "6.2"),
    72: ("6. まとめ", "6.3"),
    75: ("6. まとめ", "6.4"),
    78: ("6. まとめ", "6.5"),
}


def main():
    data = parse_markdown(DRAFT)
    cover = data["cover"]
    sections = data["sections"]

    doc = Document(str(TEMPLATE))

    # ----- Fill cover table -----
    table = doc.tables[0]
    cover_values = {
        0: cover.get("応募代表者氏名", ""),
        1: cover.get("チーム名", ""),
        2: cover.get("選択課題ID", ""),
        3: cover.get("選択課題名", ""),
        4: cover.get("提案名", ""),
        5: "",  # 解決案ID - leave empty or fill if available
    }
    for row_idx, value in cover_values.items():
        cell = table.rows[row_idx].cells[1]
        # Clear existing text
        for p in cell.paragraphs:
            clear_paragraph(p)
        if cell.paragraphs:
            cell.paragraphs[0].add_run(value)

    # ----- Fill sections -----
    # IMPORTANT: Capture paragraph references BEFORE any insertions,
    # then process in reverse index order so that insertions at higher
    # indices don't shift the positions of lower-index paragraphs.
    paragraphs = list(doc.paragraphs)  # snapshot the list

    # Also capture references for AI section and "以上" before modification
    ai_section = None
    for sk in sections:
        if "AI活用" in sk:
            ai_section = sk
            break

    ai_insert_ref = paragraphs[79] if len(paragraphs) > 79 else None

    # Build list of (para_ref, content) to process, sorted by index descending
    tasks = []
    for para_idx in sorted(SECTION_MAP.keys(), reverse=True):
        section_key, subsection = SECTION_MAP[para_idx]
        if para_idx >= len(paragraphs):
            print(f"Warning: paragraph index {para_idx} out of range")
            continue

        matched_section = None
        for sk in sections:
            if section_key in sk or sk in section_key:
                matched_section = sk
                break

        if matched_section is None:
            print(f"Warning: section '{section_key}' not found in markdown")
            continue

        section_text = sections[matched_section]

        if subsection == "__FULL__":
            content = get_section_body(section_text, include_subsections=True)
        elif subsection:
            content = extract_subsection(section_text, subsection)
        else:
            content = get_section_body(section_text)

        if not content:
            print(f"Warning: no content for {section_key} / {subsection}")
            continue

        tasks.append((paragraphs[para_idx], content))

    # Process in reverse order (highest index first)
    for placeholder_para, content in tasks:
        clear_paragraph(placeholder_para)
        add_formatted_content(doc, placeholder_para, content)

    # ----- Handle AI活用方針 -----
    if ai_section and ai_insert_ref is not None:
        heading_p = add_paragraph_after(doc, ai_insert_ref)
        run = heading_p.add_run("AI活用方針")
        run.font.bold = True

        ai_text = sections[ai_section]
        ai_text = re.sub(r"^---\s*$", "", ai_text, flags=re.MULTILINE).strip()
        ai_text = ai_text.replace("以上", "").strip()
        add_formatted_content(doc, heading_p, ai_text)

    # ----- Save -----
    doc.save(str(OUTPUT))
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()
