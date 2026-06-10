#!/usr/bin/env python3
"""
Generate report_official.docx from the OFFICIAL NEDO template
(report_template_official.docx) + report_draft.md.

- Fills the 6-row cover table (応募代表者氏名 / チーム名 / 選択課題ID /
  選択課題名 / 提案名 / 解決案ID). The name/team are PROVISIONAL pending
  事務局 confirmation of the anonymity rule; blank them by editing COVER below.
- Replaces each section's blue guidance paragraph with the markdown content.
- Maps draft headings onto the official template's fixed headings
  (e.g. draft 2.3 -> template "2.3 古典コンピュータでの現状の認識").
- Appends extra draft material (6.6 今後の展望, AI活用方針) before "以上".
- Body font: MS PGothic 11pt (per template 留意点).
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING

BASE_DIR = Path(__file__).parent
TEMPLATE = BASE_DIR / "report_template_official.docx"
DRAFT = BASE_DIR / "report_draft.md"
OUTPUT = BASE_DIR / "report_official.docx"

JP_FONT = "ＭＳ Ｐゴシック"
BODY_PT = Pt(11)

# Cover values. Set NAME/TEAM to "" to anonymize after 事務局 confirmation.
COVER_FROM_MD = True  # pull values from the markdown 表紙情報 block


# ---------------------------------------------------------------------------
# Parse the markdown
# ---------------------------------------------------------------------------

def parse_markdown(path: Path) -> dict:
    text = path.read_text(encoding="utf-8")
    lines = text.split("\n")

    cover = {}
    cover_keys = ["応募代表者氏名", "チーム名", "選択課題ID",
                  "選択課題名", "提案名", "解決案ID"]
    for line in lines:
        for key in cover_keys:
            if line.strip().startswith(f"- {key}:"):
                cover[key] = line.split(":", 1)[1].strip()

    sections = {}
    current_key, current_lines = None, []
    for line in lines:
        m = re.match(r"^##\s+(.+)", line)
        if m:
            if current_key is not None:
                sections[current_key] = "\n".join(current_lines).strip()
            current_key, current_lines = m.group(1).strip(), []
        elif current_key is not None:
            current_lines.append(line)
    if current_key is not None:
        sections[current_key] = "\n".join(current_lines).strip()

    return {"cover": cover, "sections": sections}


def extract_subsection(section_text: str, prefix: str) -> str:
    """Content under a ### subsection that starts with `prefix` (e.g. '2.3')."""
    lines = section_text.split("\n")
    collecting, result = False, []
    for line in lines:
        if re.match(r"^###\s+" + re.escape(prefix) + r"\b", line):
            collecting = True
            continue
        if re.match(r"^###\s+", line):
            if collecting:
                break
            continue
        if collecting:
            result.append(line)
    return "\n".join(result).strip()


def section_full_body(section_text: str) -> str:
    """Whole section incl. ### subsections, with --- rules stripped."""
    return re.sub(r"^---\s*$", "", section_text, flags=re.MULTILINE).strip()


# ---------------------------------------------------------------------------
# docx helpers
# ---------------------------------------------------------------------------

def _apply_jp(run, size=BODY_PT, bold=False, mono=False, color=None):
    name = "Courier New" if mono else JP_FONT
    run.font.name = name
    run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = run._element.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)


def clear_paragraph(para):
    el = para._element
    for child in list(el):
        if child.tag == qn("w:r"):
            el.remove(child)


def add_paragraph_after(doc, ref_para, text="", style=None):
    new_para = doc.add_paragraph(text, style=style)
    # Explicit, compact line spacing. The template docDefault uses lineRule="auto"
    # which, with CJK font metrics (large line height), renders very loose (~1.6-1.8x).
    # Fix to an exact line height for a normal, professional density.
    pf = new_para.paragraph_format
    pf.line_spacing = Pt(15)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.space_after = Pt(4)
    pf.space_before = Pt(0)
    ref_para._element.addnext(new_para._element)
    return new_para


def _add_rich_text(para, text):
    """Handle **bold** and `inline code`; rest plain. All MS PGothic 11pt."""
    tokens = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = para.add_run(tok[2:-2]); _apply_jp(r, bold=True)
        elif tok.startswith("`") and tok.endswith("`"):
            r = para.add_run(tok[1:-1]); _apply_jp(r, mono=True)
        else:
            r = para.add_run(tok); _apply_jp(r)


def add_formatted_content(doc, ref_para, content_text):
    lines = content_text.split("\n")
    last = ref_para
    i = 0
    while i < len(lines):
        line = lines[i]

        # fenced code block
        if line.strip().startswith("```"):
            i += 1
            code = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i]); i += 1
            i += 1
            p = add_paragraph_after(doc, last)
            r = p.add_run("\n".join(code))
            _apply_jp(r, size=Pt(8), mono=True, color=RGBColor(0x33, 0x33, 0x33))
            pPr = p._element.get_or_add_pPr()
            shd = p._element.makeelement(qn("w:shd"), {
                qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "F2F2F2"})
            pPr.append(shd)
            last = p
            continue

        # markdown table
        if line.strip().startswith("|") and i + 1 < len(lines) and \
                set(lines[i + 1].strip().replace("|", "").replace(":", "").strip()) <= {"-", " "} and \
                lines[i + 1].strip().startswith("|"):
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i]); i += 1
            header = [c.strip() for c in tbl_lines[0].strip().strip("|").split("|")]
            rows = [[c.strip() for c in tl.strip().strip("|").split("|")] for tl in tbl_lines[2:]]
            ncols = len(header)
            table = doc.add_table(rows=len(rows) + 1, cols=ncols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            for ci, val in enumerate(header):
                cell = table.rows[0].cells[ci]
                clear_paragraph(cell.paragraphs[0])
                _add_table_cell(cell.paragraphs[0], val, bold=True)
            for ri, rdata in enumerate(rows):
                for ci in range(min(len(rdata), ncols)):
                    cell = table.rows[ri + 1].cells[ci]
                    clear_paragraph(cell.paragraphs[0])
                    _add_table_cell(cell.paragraphs[0], rdata[ci])
            last._element.addnext(table._element)
            spacer = add_paragraph_after(doc, last)
            spacer._element.addprevious(table._element)
            last = spacer
            continue

        if line.strip() in ("", "---"):
            i += 1
            continue

        # markdown subheading (### / #### ...) -> bold subheading
        m_h = re.match(r"^(#{2,6})\s+(.*)", line)
        if m_h:
            level = len(m_h.group(1))
            p = add_paragraph_after(doc, last)
            r = p.add_run(m_h.group(2).strip())
            _apply_jp(r, size=Pt(12 if level <= 3 else 11), bold=True)
            last = p
            i += 1
            continue

        # bullet list item
        m_b = re.match(r"^\s*[-*]\s+(.*)", line)
        if m_b:
            p = add_paragraph_after(doc, last, style="List Paragraph")
            p.paragraph_format.left_indent = Pt(18)
            r = p.add_run("• "); _apply_jp(r)
            _add_rich_text(p, m_b.group(1))
            last = p
            i += 1
            continue

        # plain paragraph: merge consecutive soft-wrapped lines into ONE paragraph.
        # report_draft.md hard-wraps mid-sentence; emitting one paragraph per line made
        # the template's 8pt paragraph-after spacing apply between every wrapped line
        # (huge line gaps + page bloat). Join with "" since JP wraps without spaces.
        buf = [line.strip()]
        i += 1
        while i < len(lines):
            nl = lines[i]; nls = nl.strip()
            if nls in ("", "---") or nls.startswith("```") or nls.startswith("|"):
                break
            if re.match(r"^(#{2,6})\s+", nl) or re.match(r"^\s*[-*]\s+", nl):
                break
            buf.append(nls); i += 1
        p = add_paragraph_after(doc, last)
        _add_rich_text(p, "".join(buf))
        last = p
    return last


def _add_table_cell(para, text, bold=False):
    tokens = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = para.add_run(tok[2:-2]); _apply_jp(r, size=Pt(9), bold=True)
        elif tok.startswith("`") and tok.endswith("`"):
            r = para.add_run(tok[1:-1]); _apply_jp(r, size=Pt(9), mono=True)
        else:
            r = para.add_run(tok); _apply_jp(r, size=Pt(9), bold=bold)


# ---------------------------------------------------------------------------
# Mapping: template section-number -> (md ## key, ### prefix or marker)
# ---------------------------------------------------------------------------

S2 = "2. 課題の内容・研究の目的"
S3 = "3. 解決案の内容"
S5 = "5. 考察"
S6 = "6. まとめ"

NUM_MAP = {
    "1":   ("1. 要約（1000文字以内）", "__BODY__"),
    "2.1": (S2, "2.1"), "2.2": (S2, "2.2"), "2.3": (S2, "2.3"), "2.4": (S2, "2.4"),
    "3.1": (S3, "3.1"), "3.2": (S3, "3.2"), "3.3": (S3, "3.3"),
    "3.4": (S3, "3.4"), "3.5": (S3, "3.5"),
    "4":   ("4. 研究結果", "__FULL__"),
    "5.1": (S5, "5.1"), "5.2": (S5, "5.2"), "5.3": (S5, "5.3"), "5.4": (S5, "5.4"),
    "6.1": (S6, "6.1"), "6.2": (S6, "6.2"), "6.3": (S6, "6.3"),
    "6.4": (S6, "6.4"), "6.5": (S6, "6.5"),
}


def heading_number(text: str):
    """Return the leading section number of a heading paragraph, or None."""
    t = text.strip().replace("　", " ")
    m = re.match(r"^(\d+(?:\.\d+)?)[\.\s　]", t + " ")
    return m.group(1) if m else None


def resolve_content(sections, md_key, marker):
    if md_key not in sections:
        # tolerant match
        for k in sections:
            if md_key.split("（")[0][:4] in k:
                md_key = k
                break
    text = sections.get(md_key, "")
    if marker == "__BODY__":
        body = re.sub(r"^---\s*$", "", text, flags=re.MULTILINE)
        out = []
        for ln in body.split("\n"):
            if re.match(r"^###\s+", ln):
                break
            out.append(ln)
        return "\n".join(out).strip()
    if marker == "__FULL__":
        return section_full_body(text)
    return extract_subsection(text, marker)


def _localname(tag):
    return tag.split("}")[-1]


def remove_template_furniture(doc):
    """Remove the 【記入にあたっての留意点】 instruction box and the template
    version footer — these are template instructions, not report content."""
    body = doc.element.body
    runs_to_remove = []
    for tx in body.iter():
        if _localname(tx.tag) == "txbxContent":
            txt = "".join(n.text or "" for n in tx.iter(qn("w:t")))
            if "留意点" in txt or "別の書式での提出は認められません" in txt:
                anc = tx
                while anc is not None and _localname(anc.tag) != "r":
                    anc = anc.getparent()
                if anc is not None and anc not in runs_to_remove:
                    runs_to_remove.append(anc)
    for r in runs_to_remove:
        r.getparent().remove(r)
    # clear version footer ("Ver. ....")
    for sec in doc.sections:
        for p in sec.footer.paragraphs:
            if p.text.strip().startswith("Ver."):
                clear_paragraph(p)


def remove_blank_lead(doc):
    """Remove leftover empty paragraphs between the cover table and the first
    content paragraph, so the report does not start with a blank page 2.
    (The template's 留意点 instructions sat here; removing them left empty shells
    that, combined with the page break before §1, pushed content to page 3.)"""
    body = doc.element.body
    cover_seen = False
    to_remove = []
    for el in body.iterchildren():
        tag = el.tag.split("}")[-1]
        if tag == "tbl":
            txt = "".join(t.text or "" for t in el.iter(qn("w:t")))
            if "応募代表者氏名" in txt:
                cover_seen = True
            continue
        if tag != "p" or not cover_seen:
            continue
        txt = "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()
        has_break = any(b.get(qn("w:type")) == "page" for b in el.iter(qn("w:br")))
        has_sect = el.find(".//" + qn("w:sectPr")) is not None
        if txt == "" and not has_break and not has_sect:
            to_remove.append(el)
        else:
            break
    for el in to_remove:
        el.getparent().remove(el)


def main():
    data = parse_markdown(DRAFT)
    cover, sections = data["cover"], data["sections"]
    doc = Document(str(TEMPLATE))
    remove_template_furniture(doc)

    # set Normal style default font
    normal = doc.styles["Normal"]
    normal.font.name = JP_FONT
    normal.font.size = BODY_PT
    rpr = normal.element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = rpr.makeelement(qn("w:rFonts"), {}); rpr.insert(0, rf)
    for a in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        rf.set(qn(a), JP_FONT)

    # ----- cover table -----
    table = doc.tables[0]
    order = ["応募代表者氏名", "チーム名", "選択課題ID", "選択課題名", "提案名", "解決案ID"]
    for ri, key in enumerate(order):
        val = cover.get(key, "") if COVER_FROM_MD else ""
        cell = table.rows[ri].cells[1]
        for p in cell.paragraphs:
            clear_paragraph(p)
        r = cell.paragraphs[0].add_run(val)
        _apply_jp(r)

    # ----- locate guidance paragraphs to replace -----
    paras = list(doc.paragraphs)
    tasks = []  # (guidance_para, content)
    closing_para = None
    for idx, p in enumerate(paras):
        if p.text.strip() == "以上":
            closing_para = p
        num = heading_number(p.text)
        if num in NUM_MAP and idx + 1 < len(paras):
            md_key, marker = NUM_MAP[num]
            content = resolve_content(sections, md_key, marker)
            if content:
                tasks.append((paras[idx + 1], content))
            else:
                print(f"WARN: no content for section {num}")

    # process bottom-up so element insertions don't disturb earlier refs
    for guidance_para, content in reversed(tasks):
        clear_paragraph(guidance_para)
        add_formatted_content(doc, guidance_para, content)

    # ----- extra material before 以上: 6.6 今後の展望 + AI活用方針 -----
    if closing_para is not None:
        extras = []
        s6 = sections.get(S6, "")
        body66 = extract_subsection(s6, "6.6")
        if body66:
            extras.append(("6.6 今後の展望", body66))
        for k in sections:
            if "AI活用" in k:
                ai = re.sub(r"^---\s*$", "", sections[k], flags=re.MULTILINE)
                ai = ai.replace("以上", "").strip()
                if ai:
                    extras.append(("AI活用方針", ai))
                break
        # anchor: empty para right before 以上
        anchor = doc.add_paragraph("")
        closing_para._element.addprevious(anchor._element)
        last = anchor
        for title, body in extras:
            h = add_paragraph_after(doc, last)
            r = h.add_run(title); _apply_jp(r, size=Pt(13), bold=True)
            last = add_formatted_content(doc, h, body)

    # Anonymize document metadata (template carries 作成者 'Hiroki Okuda (JP)').
    remove_blank_lead(doc)

    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""

    doc.save(str(OUTPUT))
    print(f"Saved: {OUTPUT}")
    print(f"Filled sections: {len(tasks)}")


if __name__ == "__main__":
    main()
